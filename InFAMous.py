import inspect
import io
import os
import re
import sqlite3
import tempfile
import uuid
from contextlib import closing
from datetime import datetime
from typing import List, Optional

import streamlit as st
from dotenv import load_dotenv
from docx import Document
from google import genai
from google.genai import types
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

# =========================
# CONFIG
# =========================
load_dotenv()

API_KEY = os.getenv("GEMINI_API_KEY")
DB_FILE = "agent_x.db"
CHAT_MODEL = "gemini-2.5-flash"
EXPORT_TITLE = "Agent X Output"

if not API_KEY:
    st.error("GEMINI_API_KEY not found in your .env file.")
    st.stop()

client = genai.Client(api_key=API_KEY)

st.set_page_config(page_title="Agent X", page_icon="🤖", layout="wide")

# =========================
# STYLES
# =========================
st.markdown(
    """
    <style>
    .stApp {
        background-image:
            linear-gradient(rgba(0,0,0,0.84), rgba(0,0,0,0.84)),
            url("https://cdn.mos.cms.futurecdn.net/3QnBSmh9o3uy38kz7HfBZC-1920-80.jpg.webp");
        background-size: cover;
        background-position: center;
        background-repeat: no-repeat;
        background-attachment: fixed;
    }

    [data-testid="stHeader"] {
        background: rgba(0,0,0,0);
    }

    [data-testid="stSidebar"] {
        background: rgba(10, 14, 18, 0.94);
        border-right: 1px solid rgba(255,255,255,0.06);
    }

    .block-container {
        padding-top: 1.3rem;
        padding-bottom: 6rem;
        max-width: none !important;
        width: 100%;
    }

    h1, h2, h3, h4, p, div, span, label {
        color: white !important;
    }

    .app-title {
        font-size: 34px;
        font-weight: 700;
        margin-bottom: 2px;
    }

    .app-subtitle {
        font-size: 14px;
        opacity: 0.78;
        margin-bottom: 10px;
    }

    .top-caption {
        font-size: 13px;
        opacity: 0.72;
        margin-bottom: 18px;
    }

    .meta-card {
        background: rgba(255,255,255,0.03);
        border: 1px solid rgba(255,255,255,0.06);
        border-radius: 12px;
        padding: 8px 10px;
        margin-bottom: 8px;
    }

    .meta-title {
        font-size: 14px;
        font-weight: 600;
        color: white !important;
        margin-bottom: 2px;
    }

    .meta-subtitle {
        font-size: 11px;
        opacity: 0.6;
        color: white !important;
    }

    .sidebar-section-title {
        font-size: 12px;
        text-transform: uppercase;
        letter-spacing: 0.6px;
        opacity: 0.72;
        margin-top: 10px;
        margin-bottom: 8px;
        font-weight: 700;
    }

    .stButton button, .stDownloadButton button {
        border-radius: 12px !important;
        width: 100%;
    }

    .voice-note {
        font-size: 12px;
        opacity: 0.72;
        margin-bottom: 8px;
    }

    .source-box {
        max-width: 760px;
        margin: 6px 0 14px 0;
        padding: 8px 12px;
        background: rgba(255,255,255,0.04);
        border-radius: 12px;
        border: 1px solid rgba(255,255,255,0.05);
    }

    .source-title {
        font-size: 11px;
        opacity: 0.7;
        margin-bottom: 5px;
        font-weight: 700;
    }

    .grounded-pill {
        display: inline-block;
        font-size: 11px;
        opacity: 0.92;
        padding: 4px 8px;
        margin: 6px 0 8px 0;
        border-radius: 999px;
        background: rgba(74, 163, 255, 0.18);
        border: 1px solid rgba(74, 163, 255, 0.28);
        color: white !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# =========================
# SYSTEM PROMPT
# =========================
SYSTEM_PROMPT = """

You are Agent X, a helpful, sharp,and a useful AI assistant.

Identity rules:
- Your name is Agent X.
- Never say your name is Gemini, Google AI, Google Assistant, or anything else.
- If the user asks who you are, say: "I'm Agent X."
- Speak as Agent X at all times.

Behavior rules:
- Be clear, helpful, and educational.
- Be a little witty when appropriate, but not offensive.
- If the user uploads an image, describe what is visible and explain it clearly.
- If the user uploads a document or PDF, answer from the file when possible.
- If the file does not contain the answer, say that clearly.
- Format strong answers cleanly with short sections when useful.
- If the user asks you to generate an image and you ran out of quota, say try again later rather than showing a raw error.
- For factual, recent, news, sports, celebrity, product, or current-event questions, prefer live search grounding.
- If the user sends audio, transcribe the audio first internally, then answer the user naturally.
"""

# =========================
# DATABASE
# =========================
def get_conn():
    conn = sqlite3.connect(DB_FILE, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def now_str():
    return datetime.now().strftime("%b %d, %I:%M %p")


def init_db():
    conn = get_conn()
    with closing(conn.cursor()) as cur:
        cur.execute("""
            CREATE TABLE IF NOT EXISTS projects (
                id TEXT PRIMARY KEY,
                title TEXT NOT NULL,
                created_at TEXT NOT NULL
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS chats (
                id TEXT PRIMARY KEY,
                project_id TEXT NOT NULL,
                title TEXT NOT NULL,
                created_at TEXT NOT NULL,
                FOREIGN KEY (project_id) REFERENCES projects (id)
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS messages (
                id TEXT PRIMARY KEY,
                chat_id TEXT NOT NULL,
                role TEXT NOT NULL,
                content TEXT NOT NULL,
                created_at TEXT NOT NULL,
                FOREIGN KEY (chat_id) REFERENCES chats (id)
            )
        """)
        conn.commit()

    with closing(conn.cursor()) as cur:
        cur.execute("SELECT COUNT(*) AS count FROM projects")
        count = cur.fetchone()["count"]
        if count == 0:
            project_id = str(uuid.uuid4())
            chat_id = str(uuid.uuid4())
            now = now_str()

            cur.execute(
                "INSERT INTO projects (id, title, created_at) VALUES (?, ?, ?)",
                (project_id, "New Project", now),
            )
            cur.execute(
                "INSERT INTO chats (id, project_id, title, created_at) VALUES (?, ?, ?, ?)",
                (chat_id, project_id, "New Chat", now),
            )
            conn.commit()

    conn.close()


def list_projects():
    conn = get_conn()
    rows = conn.execute("SELECT * FROM projects ORDER BY rowid DESC").fetchall()
    conn.close()
    return rows


def list_chats(project_id: str):
    conn = get_conn()
    rows = conn.execute(
        "SELECT * FROM chats WHERE project_id = ? ORDER BY rowid DESC",
        (project_id,),
    ).fetchall()
    conn.close()
    return rows


def list_messages(chat_id: str):
    conn = get_conn()
    rows = conn.execute(
        "SELECT * FROM messages WHERE chat_id = ? ORDER BY rowid ASC",
        (chat_id,),
    ).fetchall()
    conn.close()
    return rows


def create_project(title="New Project"):
    conn = get_conn()
    project_id = str(uuid.uuid4())
    chat_id = str(uuid.uuid4())
    now = now_str()

    conn.execute(
        "INSERT INTO projects (id, title, created_at) VALUES (?, ?, ?)",
        (project_id, title, now),
    )
    conn.execute(
        "INSERT INTO chats (id, project_id, title, created_at) VALUES (?, ?, ?, ?)",
        (chat_id, project_id, "New Chat", now),
    )
    conn.commit()
    conn.close()
    return project_id, chat_id


def create_chat(project_id: str, title="New Chat"):
    conn = get_conn()
    chat_id = str(uuid.uuid4())
    now = now_str()
    conn.execute(
        "INSERT INTO chats (id, project_id, title, created_at) VALUES (?, ?, ?, ?)",
        (chat_id, project_id, title, now),
    )
    conn.commit()
    conn.close()
    return chat_id


def rename_project(project_id: str, new_title: str):
    conn = get_conn()
    conn.execute(
        "UPDATE projects SET title = ? WHERE id = ?",
        (new_title.strip() or "Untitled Project", project_id),
    )
    conn.commit()
    conn.close()


def rename_chat(chat_id: str, new_title: str):
    conn = get_conn()
    conn.execute(
        "UPDATE chats SET title = ? WHERE id = ?",
        (new_title.strip() or "Untitled Chat", chat_id),
    )
    conn.commit()
    conn.close()


def delete_project(project_id: str):
    conn = get_conn()
    chat_rows = conn.execute(
        "SELECT id FROM chats WHERE project_id = ?", (project_id,)
    ).fetchall()
    for row in chat_rows:
        conn.execute("DELETE FROM messages WHERE chat_id = ?", (row["id"],))
    conn.execute("DELETE FROM chats WHERE project_id = ?", (project_id,))
    conn.execute("DELETE FROM projects WHERE id = ?", (project_id,))
    conn.commit()
    conn.close()


def delete_chat(chat_id: str):
    conn = get_conn()
    conn.execute("DELETE FROM messages WHERE chat_id = ?", (chat_id,))
    conn.execute("DELETE FROM chats WHERE id = ?", (chat_id,))
    conn.commit()
    conn.close()


def add_message(chat_id: str, role: str, content: str):
    conn = get_conn()
    conn.execute(
        "INSERT INTO messages (id, chat_id, role, content, created_at) VALUES (?, ?, ?, ?, ?)",
        (str(uuid.uuid4()), chat_id, role, content, now_str()),
    )
    conn.commit()
    conn.close()


def update_chat_title_if_new(chat_id: str, title: str):
    conn = get_conn()
    row = conn.execute("SELECT title FROM chats WHERE id = ?", (chat_id,)).fetchone()
    if row and row["title"] == "New Chat":
        conn.execute(
            "UPDATE chats SET title = ? WHERE id = ?",
            (title[:32] + ("..." if len(title) > 32 else ""), chat_id),
        )
        conn.commit()
    conn.close()


def get_project(project_id: str):
    conn = get_conn()
    row = conn.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    conn.close()
    return row


def get_chat(chat_id: str):
    conn = get_conn()
    row = conn.execute("SELECT * FROM chats WHERE id = ?", (chat_id,)).fetchone()
    conn.close()
    return row


def get_latest_assistant_message(chat_id: str) -> Optional[str]:
    conn = get_conn()
    row = conn.execute(
        """
        SELECT content FROM messages
        WHERE chat_id = ? AND role = 'assistant'
        ORDER BY rowid DESC
        LIMIT 1
        """,
        (chat_id,),
    ).fetchone()
    conn.close()
    return row["content"] if row else None


# =========================
# HELPERS
# =========================
def save_uploaded_file_temporarily(uploaded_file) -> str:
    suffix = os.path.splitext(uploaded_file.name)[1]
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.getbuffer())
        return tmp.name


def upload_files_to_gemini(uploaded_files) -> List[object]:
    uploaded_refs = []
    temp_paths = []

    try:
        for uploaded_file in uploaded_files:
            path = save_uploaded_file_temporarily(uploaded_file)
            temp_paths.append(path)
            file_ref = client.files.upload(file=path)
            uploaded_refs.append(file_ref)
    finally:
        for path in temp_paths:
            try:
                os.remove(path)
            except Exception:
                pass

    return uploaded_refs


def build_history_text(messages) -> str:
    lines = [SYSTEM_PROMPT.strip(), "", "Conversation:"]
    for m in messages:
        role = "User" if m["role"] == "user" else "Assistant"
        lines.append(f"{role}: {m['content']}")
    lines.append("Assistant:")
    return "\n".join(lines)


def extract_response_text(response) -> str:
    if getattr(response, "text", None):
        return response.text

    texts = []
    if getattr(response, "candidates", None):
        for candidate in response.candidates:
            content = getattr(candidate, "content", None)
            if content and getattr(content, "parts", None):
                for part in content.parts:
                    if getattr(part, "text", None):
                        texts.append(part.text)

    return "\n".join(texts).strip() if texts else "No response returned."


def extract_sources(response) -> List[dict]:
    sources = []

    try:
        candidates = getattr(response, "candidates", None) or []
        for candidate in candidates:
            grounding_metadata = getattr(candidate, "grounding_metadata", None)
            if not grounding_metadata:
                continue

            grounding_chunks = getattr(grounding_metadata, "grounding_chunks", None) or []
            for chunk in grounding_chunks:
                web_info = getattr(chunk, "web", None)
                if not web_info:
                    continue

                title = getattr(web_info, "title", None) or "Untitled source"
                uri = getattr(web_info, "uri", None) or ""

                if uri and not any(s["url"] == uri for s in sources):
                    sources.append({"title": title, "url": uri})
    except Exception:
        pass

    return sources


def render_sources(sources: List[dict]):
    if not sources:
        return

    links_html = ""
    for src in sources:
        title = (
            src.get("title", "Untitled source")
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
        url = src.get("url", "")
        links_html += f'<div><a href="{url}" target="_blank">{title}</a></div>'

    st.markdown(
        f"""
        <div class="source-box">
            <div class="source-title">Sources</div>
            {links_html}
        </div>
        """,
        unsafe_allow_html=True,
    )


def generate_grounded_reply(contents, user_prompt: str):
    search_tool = types.Tool(google_search=types.GoogleSearch())

    response = client.models.generate_content(
        model=CHAT_MODEL,
        contents=contents,
        config=types.GenerateContentConfig(
            tools=[search_tool],
            system_instruction=(
                "For factual, recent, news, sports, celebrity, product, or current-event questions, "
                "use Google Search grounding and answer using grounded web results when available."
            ),
        ),
    )

    reply = extract_response_text(response)
    sources = extract_sources(response)

    if not sources and user_prompt.strip():
        retry_contents = list(contents)
        retry_contents.append(
            f"IMPORTANT: Use Google Search grounding for this question and answer with fresh web results if available.\n\nQuestion: {user_prompt}"
        )

        response = client.models.generate_content(
            model=CHAT_MODEL,
            contents=retry_contents,
            config=types.GenerateContentConfig(
                tools=[search_tool],
                system_instruction=(
                    "You must prefer Google Search grounding for factual or current questions. "
                    "If search results are available, answer using them."
                ),
            ),
        )

        reply = extract_response_text(response)
        sources = extract_sources(response)

    return reply, sources


def make_pdf_bytes(title: str, content: str) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=LETTER)
    styles = getSampleStyleSheet()

    story = [Paragraph(title, styles["Title"]), Spacer(1, 12)]
    for block in content.split("\n\n"):
        block = block.strip()
        if block:
            story.append(Paragraph(block.replace("\n", "<br/>"), styles["BodyText"]))
            story.append(Spacer(1, 8))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def make_docx_bytes(title: str, content: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for block in content.split("\n\n"):
        block = block.strip()
        if block:
            doc.add_paragraph(block)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def wants_pdf(prompt: str) -> bool:
    prompt = (prompt or "").lower().strip()
    return any(
        phrase in prompt
        for phrase in [
            "make pdf", "as pdf", "export pdf", "download pdf",
            "generate pdf", "save as pdf", "create pdf"
        ]
    )


def wants_docx(prompt: str) -> bool:
    prompt = (prompt or "").lower().strip()
    return any(
        phrase in prompt
        for phrase in [
            "make docx", "as docx", "export docx", "download docx",
            "generate docx", "save as docx", "create docx",
            "make word", "as word", "export word", "download word",
            "generate word", "save as word", "create word"
        ]
    )


def clean_export_request(prompt: str) -> str:
    cleaned = prompt or ""
    patterns = [
        r"\bmake pdf\b", r"\bas pdf\b", r"\bexport pdf\b", r"\bdownload pdf\b",
        r"\bgenerate pdf\b", r"\bsave as pdf\b", r"\bcreate pdf\b",
        r"\bmake docx\b", r"\bas docx\b", r"\bexport docx\b", r"\bdownload docx\b",
        r"\bgenerate docx\b", r"\bsave as docx\b", r"\bcreate docx\b",
        r"\bmake word\b", r"\bas word\b", r"\bexport word\b", r"\bdownload word\b",
        r"\bgenerate word\b", r"\bsave as word\b", r"\bcreate word\b",
    ]
    for pattern in patterns:
        cleaned = re.sub(pattern, "", cleaned, flags=re.IGNORECASE)
    return " ".join(cleaned.split()).strip()


def auto_chat_title_from_prompt(prompt: str) -> str:
    text = (prompt or "").strip()
    if not text or text == "(sent an attachment)":
        return "New Chat"
    return text[:32] + ("..." if len(text) > 32 else "")


def render_message(role: str, content: str):
    label = "You" if role == "user" else "Agent X"
    bubble_bg = "rgba(35, 47, 66, 0.95)" if role == "user" else "rgba(17, 23, 33, 0.88)"
    bubble_radius = "16px 16px 6px 16px" if role == "user" else "16px 16px 16px 6px"

    safe_content = (
        content.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br>")
    )

    justify = "flex-end" if role == "user" else "flex-start"

    st.markdown(
        f"""
        <div style="
            display:flex;
            justify-content:{justify};
            width:100%;
            margin:10px 0;
        ">
            <div style="
                display:inline-block;
                width:auto;
                max-width:760px;
                background:{bubble_bg};
                padding:12px 15px;
                border-radius:{bubble_radius};
                font-size:15px;
                line-height:1.55;
                box-shadow:0 8px 24px rgba(0,0,0,0.18);
            ">
                <div style="font-size:11px;opacity:.68;margin-bottom:5px;font-weight:700;">
                    {label}
                </div>
                <div>{safe_content}</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def build_chat_input():
    return st.chat_input(
        placeholder="Message Agent X or attach files...",
        accept_file=True,
        file_type=["png", "jpg", "jpeg", "webp", "pdf", "txt", "md"],
        key="main_chat_input_box",
    )


# =========================
# APP INIT
# =========================
init_db()

if "active_project_id" not in st.session_state:
    projects = list_projects()
    st.session_state.active_project_id = projects[0]["id"]

if "active_chat_id" not in st.session_state:
    chats = list_chats(st.session_state.active_project_id)
    st.session_state.active_chat_id = chats[0]["id"]

if "rename_project_mode" not in st.session_state:
    st.session_state.rename_project_mode = None

if "rename_chat_mode" not in st.session_state:
    st.session_state.rename_chat_mode = None

if "pending_pdf" not in st.session_state:
    st.session_state.pending_pdf = None

if "pending_docx" not in st.session_state:
    st.session_state.pending_docx = None

if "last_sources" not in st.session_state:
    st.session_state.last_sources = []


def ensure_valid_active_ids():
    projects = list_projects()
    if not projects:
        project_id, chat_id = create_project()
        st.session_state.active_project_id = project_id
        st.session_state.active_chat_id = chat_id
        return

    project_ids = [p["id"] for p in projects]
    if st.session_state.active_project_id not in project_ids:
        st.session_state.active_project_id = projects[0]["id"]

    chats = list_chats(st.session_state.active_project_id)
    if not chats:
        chat_id = create_chat(st.session_state.active_project_id)
        st.session_state.active_chat_id = chat_id
        return

    chat_ids = [c["id"] for c in chats]
    if st.session_state.active_chat_id not in chat_ids:
        st.session_state.active_chat_id = chats[0]["id"]


ensure_valid_active_ids()

# =========================
# SIDEBAR
# =========================
with st.sidebar:
    st.subheader("Workspace")

    if st.button("➕ New Project", use_container_width=True, key="new_project_btn_main"):
        project_id, chat_id = create_project()
        st.session_state.active_project_id = project_id
        st.session_state.active_chat_id = chat_id
        st.session_state.rename_project_mode = project_id
        st.rerun()

    st.markdown('<div class="sidebar-section-title">Projects</div>', unsafe_allow_html=True)

    projects = list_projects()

    for project in projects:
        is_active_project = project["id"] == st.session_state.active_project_id

        st.markdown(
            f"""
            <div class="meta-card">
                <div class="meta-title">{project['title']}</div>
                <div class="meta-subtitle">{project['created_at']}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        col1, col2, col3 = st.columns([5, 1, 1])

        with col1:
            if st.button(
                f"{'📂' if is_active_project else '📁'} Open",
                key=f"open_project_{project['id']}",
                use_container_width=True,
            ):
                st.session_state.active_project_id = project["id"]
                chats = list_chats(project["id"])
                st.session_state.active_chat_id = chats[0]["id"]
                st.rerun()

        with col2:
            if st.button("✏️", key=f"rename_project_{project['id']}"):
                st.session_state.rename_project_mode = project["id"]
                st.rerun()

        with col3:
            if len(projects) > 1:
                if st.button("🗑️", key=f"delete_project_{project['id']}"):
                    delete_project(project["id"])
                    ensure_valid_active_ids()
                    st.session_state.rename_project_mode = None
                    st.session_state.rename_chat_mode = None
                    st.rerun()

        if st.session_state.rename_project_mode == project["id"]:
            new_project_name = st.text_input(
                "Rename project",
                value=project["title"],
                key=f"project_name_input_{project['id']}",
            )
            save_col, cancel_col = st.columns(2)
            with save_col:
                if st.button("Save", key=f"save_project_name_{project['id']}", use_container_width=True):
                    rename_project(project["id"], new_project_name)
                    st.session_state.rename_project_mode = None
                    st.rerun()
            with cancel_col:
                if st.button("Cancel", key=f"cancel_project_name_{project['id']}", use_container_width=True):
                    st.session_state.rename_project_mode = None
                    st.rerun()

        if is_active_project:
            st.markdown('<div class="sidebar-section-title">Chats</div>', unsafe_allow_html=True)

            if st.button("➕ New Chat", key=f"new_chat_{project['id']}", use_container_width=True):
                chat_id = create_chat(project["id"])
                st.session_state.active_chat_id = chat_id
                st.session_state.rename_chat_mode = chat_id
                st.rerun()

            chats = list_chats(project["id"])
            for chat in chats:
                is_active_chat = chat["id"] == st.session_state.active_chat_id

                c1, c2, c3 = st.columns([5, 1, 1])

                with c1:
                    if st.button(
                        f"{'💬' if is_active_chat else '🗨️'} {chat['title']}",
                        key=f"open_chat_{chat['id']}",
                        use_container_width=True,
                    ):
                        st.session_state.active_chat_id = chat["id"]
                        st.rerun()

                with c2:
                    if st.button("✏️", key=f"rename_chat_{chat['id']}"):
                        st.session_state.rename_chat_mode = chat["id"]
                        st.rerun()

                with c3:
                    if len(chats) > 1:
                        if st.button("🗑️", key=f"delete_chat_{chat['id']}"):
                            delete_chat(chat["id"])
                            ensure_valid_active_ids()
                            st.session_state.rename_chat_mode = None
                            st.rerun()

                if st.session_state.rename_chat_mode == chat["id"]:
                    new_chat_name = st.text_input(
                        "Rename chat",
                        value=chat["title"],
                        key=f"chat_name_input_{chat['id']}",
                    )
                    save_chat_col, cancel_chat_col = st.columns(2)
                    with save_chat_col:
                        if st.button("Save", key=f"save_chat_name_{chat['id']}", use_container_width=True):
                            rename_chat(chat["id"], new_chat_name)
                            st.session_state.rename_chat_mode = None
                            st.rerun()
                    with cancel_chat_col:
                        if st.button("Cancel", key=f"cancel_chat_name_{chat['id']}", use_container_width=True):
                            st.session_state.rename_chat_mode = None
                            st.rerun()

            st.markdown("---")

# =========================
# MAIN
# =========================
active_project = get_project(st.session_state.active_project_id)
active_chat = get_chat(st.session_state.active_chat_id)
messages = list_messages(st.session_state.active_chat_id)

st.markdown('<div class="app-title"> Agent X</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="app-subtitle">Projects, chats, uploads, exports, voice-in-chat input, and live search grounding.</div>',
    unsafe_allow_html=True,
)
st.markdown(
    f'<div class="top-caption">Project: {active_project["title"]} • Chat: {active_chat["title"]}</div>',
    unsafe_allow_html=True,
)

if "accept_audio" in inspect.signature(st.chat_input).parameters:
    st.markdown(
        '<div class="voice-note">Mic is available in the chat box below. You can type, attach files, or record audio in one place.</div>',
        unsafe_allow_html=True,
    )

for message in messages:
    render_message(message["role"], message["content"])
audio_input = st.audio_input("Record a voice message for Agent X")
chat_value = build_chat_input()

if chat_value or audio_input:
    prompt = ""
    uploaded_files = []

    if chat_value:
        prompt = chat_value.text if hasattr(chat_value, "text") else ""
        uploaded_files = list(chat_value.files) if hasattr(chat_value, "files") and chat_value.files else []

    audio_file = audio_input if audio_input else None

    if audio_file:
        uploaded_files.append(audio_file)

    if prompt.strip():
        user_text = prompt.strip()
    elif audio_file:
        user_text = "(sent a voice message)"
    elif uploaded_files:
        user_text = "(sent an attachment)"
    else:
        user_text = ""

    if user_text:
        add_message(st.session_state.active_chat_id, "user", user_text)
        update_chat_title_if_new(
            st.session_state.active_chat_id,
            auto_chat_title_from_prompt(user_text),
        )
        render_message("user", user_text)

    try:
        wants_pdf_export = wants_pdf(prompt)
        wants_docx_export = wants_docx(prompt)
        latest_assistant = get_latest_assistant_message(st.session_state.active_chat_id)
        st.session_state.last_sources = []

        export_only_request = (
            (wants_pdf_export or wants_docx_export)
            and not uploaded_files
            and len(clean_export_request(prompt)) < 4
        )

        if export_only_request and latest_assistant:
            reply = "Done. Your export is ready below."

            if wants_pdf_export:
                st.session_state.pending_pdf = make_pdf_bytes(EXPORT_TITLE, latest_assistant)

            if wants_docx_export:
                st.session_state.pending_docx = make_docx_bytes(EXPORT_TITLE, latest_assistant)

            add_message(st.session_state.active_chat_id, "assistant", reply)
            render_message("assistant", reply)

        else:
            db_messages = list_messages(st.session_state.active_chat_id)
            contents = []

            if uploaded_files:
                refs = upload_files_to_gemini(uploaded_files)
                contents.extend(refs)

            cleaned_prompt = clean_export_request(prompt) if prompt else prompt

            if cleaned_prompt and cleaned_prompt != prompt and db_messages:
                db_messages = db_messages[:-1] + [{
                    "role": "user",
                    "content": cleaned_prompt
                }]

            history_text = build_history_text(db_messages)

            if audio_file:
             history_text += "\n\nThe user sent an audio recording. Transcribe the audio first, understand it, and then answer naturally as Agent X."

             contents.append(history_text)

            with st.spinner("Agent X is thinking..."):
                reply, sources = generate_grounded_reply(contents, prompt)

            st.session_state.last_sources = sources
            add_message(st.session_state.active_chat_id, "assistant", reply)

            if wants_pdf_export:
                st.session_state.pending_pdf = make_pdf_bytes(EXPORT_TITLE, reply)
            else:
                st.session_state.pending_pdf = None

            if wants_docx_export:
                st.session_state.pending_docx = make_docx_bytes(EXPORT_TITLE, reply)
            else:
                st.session_state.pending_docx = None

            render_message("assistant", reply)

            if sources:
                st.markdown(
                    '<div class="grounded-pill">Live web grounded</div>',
                    unsafe_allow_html=True,
                )
                render_sources(sources)
            else:
                st.caption("Answered without live web sources.")

    except Exception as e:
        st.error(f"Agent X error: {e}")

if st.session_state.pending_pdf or st.session_state.pending_docx:
    st.divider()
    st.subheader("Your export is ready")

    col1, col2 = st.columns(2)

    with col1:
        if st.session_state.pending_pdf:
            st.download_button(
                "Download PDF",
                data=st.session_state.pending_pdf,
                file_name="agent_x_output.pdf",
                mime="application/pdf",
                use_container_width=True,
                key="download_pdf_btn",
            )

    with col2:
        if st.session_state.pending_docx:
            st.download_button(
                "Download DOCX",
                data=st.session_state.pending_docx,
                file_name="agent_x_output.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="download_docx_btn",
            )